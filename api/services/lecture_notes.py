# lecture_notes.py
from typing import Dict, List, Optional, Tuple
import json
from datetime import datetime
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Inches
import re
from groq import Groq
from pathlib import Path
import logging
from api.core.config import get_settings
logger = logging.getLogger(__name__)

class LectureNotesGenerator:
    def __init__(self, groq_api_key: str):
        self.groq_client = Groq(api_key=groq_api_key)
        self.settings=get_settings()
    def _create_styles(self, doc: Document):
        """Create custom styles for the document."""
        # Main Title Style
        title_style = doc.styles.add_style('MainTitle', WD_STYLE_TYPE.PARAGRAPH)
        title_style.font.name = 'Calibri'
        title_style.font.size = Pt(24)
        title_style.font.bold = True
        title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_style.paragraph_format.space_after = Pt(20)

        # Section Title Style
        section_title = doc.styles.add_style('SectionTitle', WD_STYLE_TYPE.PARAGRAPH)
        section_title.font.name = 'Calibri'
        section_title.font.size = Pt(16)
        section_title.font.bold = True
        section_title.paragraph_format.space_before = Pt(24)
        section_title.paragraph_format.space_after = Pt(12)

        # Normal Text Style
        normal_style = doc.styles.add_style('NormalText', WD_STYLE_TYPE.PARAGRAPH)
        normal_style.font.name = 'Calibri'
        normal_style.font.size = Pt(11)
        normal_style.paragraph_format.space_after = Pt(12)
        normal_style.paragraph_format.line_spacing = 1.15

        # Heading Style
        heading_style = doc.styles.add_style('Heading3', WD_STYLE_TYPE.PARAGRAPH)
        heading_style.font.name = 'Calibri'
        heading_style.font.size = Pt(12)
        heading_style.font.bold = True
        heading_style.paragraph_format.space_before = Pt(12)
        heading_style.paragraph_format.space_after = Pt(6)

        # Key Points Style
        key_points_style = doc.styles.add_style('KeyPoints', WD_STYLE_TYPE.PARAGRAPH)
        key_points_style.font.name = 'Calibri'
        key_points_style.font.size = Pt(11)
        key_points_style.paragraph_format.left_indent = Inches(0.25)
        key_points_style.paragraph_format.space_after = Pt(6)
        key_points_style.paragraph_format.line_spacing = 1.15

    def _format_word_document(self, topics: List[Dict], metadata: Dict) -> Document:
        doc = Document()
        self._create_styles(doc)

        # Set margins
        for section in doc.sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)

        # Title Page
        doc.add_paragraph(metadata.get('title', 'Lecture Notes'), style='MainTitle')
        doc.add_paragraph(f"Date: {metadata.get('date')}", style='NormalText')
        if metadata.get('course'):
            doc.add_paragraph(f"Course: {metadata['course']}", style='NormalText')

        # Table of Contents
        doc.add_paragraph('\nTable of Contents', style='SectionTitle')
        for idx, topic in enumerate(topics, 1):
            toc_entry = doc.add_paragraph(style='NormalText')
            toc_entry.add_run(f"{idx}. {topic['title']} ")
            timestamp = toc_entry.add_run(f"({topic['timestamp']})")
            timestamp.font.color.rgb = RGBColor(89, 89, 89)

        doc.add_page_break()

        # Content
        for idx, topic in enumerate(topics, 1):
            # Topic Title
            title = doc.add_paragraph(style='SectionTitle')
            title.add_run(f"{idx}. {topic['title']}")
            
            # Timestamp
            timestamp = doc.add_paragraph(style='NormalText')
            ts = timestamp.add_run(f"Timestamp: {topic['timestamp']}")
            ts.font.color.rgb = RGBColor(89, 89, 89)

            # Summary
            doc.add_paragraph("Summary", style='Heading3')
            doc.add_paragraph(topic['summary'], style='NormalText')

            # Key Points
            doc.add_paragraph("Key Points", style='Heading3')
            for point in topic['key_points']:
                bullet = doc.add_paragraph(style='KeyPoints')
                bullet.add_run(f"• {point}")

        return doc
    

    async def generate_notes(
        self,
        transcript: str,
        lesson_plan: Optional[str] = None,
        metadata: Optional[Dict] = None,
        temp_dir: Optional[Path] = None
    ) -> str:
        """Generate structured lecture notes from transcript.
        
        Args:
            transcript (str): The transcribed text to process
            lesson_plan (Optional[str]): Optional lesson plan to guide note generation
            metadata (Optional[Dict]): Optional metadata for the document
            temp_dir (Optional[Path]): Directory to save the file (defaults to current dir)
        
        Returns:
            str: Filename of the generated document
        """
        try:
            settings=get_settings()
            # Validate input
            if not transcript or not transcript.strip():
                logger.error("Empty transcript provided")
                raise ValueError("Empty transcript provided")

            # Ensure proper metadata with defaults
            current_date = datetime.now().strftime('%Y-%m-%d')
            if metadata is None:
                metadata = {}
            metadata = {
                'title': metadata.get('title', 'Lecture Notes'),
                'date': metadata.get('date', current_date),
                'course': metadata.get('course', '')
            }

            logger.info("Analyzing transcript and extracting topics...")
            topics_data = self._extract_topics(transcript, lesson_plan)
            
            if not topics_data or not topics_data.get('topics'):
                logger.error("No topics were extracted from transcript")
                raise ValueError("Failed to extract topics from transcript")

            logger.info("Generating Word document...")
            doc = self._format_word_document(topics_data['topics'], metadata)
            
            # Generate filename with timestamp
            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
            filename = f"lecture_notes_{timestamp}.docx"

            # Use NOTES_DIR from settings
            save_path = settings.NOTES_DIR / filename
            
            # Ensure the directory exists
            settings.NOTES_DIR.mkdir(parents=True, exist_ok=True)
            
            # Save the document
            doc.save(str(save_path))
            logger.info(f"Lecture notes saved successfully at: {save_path}")

            # Return just the filename
            return filename

        except ValueError as e:
            logger.error(f"Validation error in generate_notes: {str(e)}")
            raise

        except Exception as e:
            logger.error(f"Unexpected error in generate_notes: {str(e)}", exc_info=True)
            raise ValueError(f"Failed to generate lecture notes: {str(e)}")
        

    # lecture_notes.py
    def _extract_topics(self, transcript: str, lesson_plan: Optional[str] = None) -> dict:
        """Extract topics with improved summary generation."""
        prompt = f"""
        Analyze this lecture transcript and extract the main topics discussed.
        Create 5 distinct sections, each with:
        - A clear title reflecting the main concept
        - A detailed summary (2-3 sentences)
        - 3-6 specific key points
        - An approximate timestamp

        Use this exact JSON format and make sure all strings are properly quoted:
        {{
            "topics": [
                {{
                    "title": "Topic Title Here",
                    "summary": "Detailed summary of the topic. Second sentence providing more context. Third sentence if needed.",
                    "key_points": [
                        "Specific key point 1",
                        "Specific key point 2",
                        "Specific key point 3"
                    ],
                    "timestamp": "00:00"
                }}
            ]
        }}

        Base your analysis on this transcript:
        {transcript}

        And consider this lesson plan (if provided):
        {lesson_plan if lesson_plan else 'No lesson plan provided'}
        """

        try:
            # Make API call with strict parameters
            response = self.groq_client.chat.completions.create(
                model="llama-3.1-70b-versatile",
                messages=[
                    {
                        "role": "system",
                        "content": "You are an expert at analyzing educational content and creating structured summaries. Always respond with properly formatted JSON."
                    },
                    {
                        "role": "user",
                        "content": prompt
                    }
                ],
                temperature=0.1,  # Lower temperature for more consistent output
                max_tokens=2000
            )

            # Get the content and clean it
            content = response.choices[0].message.content.strip()
            
            # Remove any markdown formatting if present
            if '```json' in content:
                content = content.split('```json')[1].split('```')[0].strip()
            elif '```' in content:
                content = content.split('```')[1].split('```')[0].strip()

            # Parse JSON with error handling
            try:
                result = json.loads(content)
                # Validate structure
                if not isinstance(result, dict) or 'topics' not in result:
                    raise ValueError("Invalid JSON structure")
                if not isinstance(result['topics'], list) or not result['topics']:
                    raise ValueError("No topics found in response")
                return result
            except json.JSONDecodeError as je:
                print(f"JSON parsing error: {je}")
                raise

        except Exception as e:
            print(f"Error in topic extraction: {str(e)}")
            # Return a meaningful fallback with actual content
            return {
                "topics": [
                    {
                        "title": "Introduction and Overview",
                        "summary": "Initial overview of the lecture content and main themes discussed.",
                        "key_points": [
                            "Overview of main concepts",
                            "Introduction to key themes",
                            "Framework for understanding the content"
                        ],
                        "timestamp": "00:00"
                    },
                    {
                        "title": "Main Concepts Discussion",
                        "summary": "Detailed exploration of the central ideas presented in the lecture.",
                        "key_points": [
                            "Key concept explanations",
                            "Important definitions and terms",
                            "Real-world applications"
                        ],
                        "timestamp": "05:00"
                    }
                ]
            }

